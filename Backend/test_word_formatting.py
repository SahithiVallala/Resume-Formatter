"""
Test Word formatting directly
"""
import os
from docx import Document

# Test data
test_resume_data = {
    'name': 'John Doe',
    'email': 'john.doe@email.com',
    'phone': '(555) 123-4567',
    'address': '123 Main St, City, State',
    'linkedin': 'linkedin.com/in/johndoe',
    'dob': '01/15/1990',
    'experience': [
        {'title': 'Senior Software Engineer at Tech Corp', 'duration': '2020-2023', 'details': ['Led development', 'Managed team']},
        {'title': 'Software Developer at StartupXYZ', 'duration': '2018-2020', 'details': ['Developed apps', 'CI/CD']}
    ],
    'education': [
        {'degree': 'Master of Science in Computer Science', 'year': '2018'},
        {'degree': 'Bachelor of Science in Software Engineering', 'year': '2016'}
    ],
    'skills': ['Python', 'JavaScript', 'React', 'Node.js', 'AWS'],
    'sections': {
        'experience': [
            'Senior Software Engineer at Tech Corp (2020-2023)',
            '- Led development of cloud-based applications',
            '- Managed team of 5 developers',
            'Software Developer at StartupXYZ (2018-2020)',
            '- Developed mobile applications'
        ],
        'education': [
            'Master of Science in Computer Science, University of Technology, 2018',
            'Bachelor of Science in Software Engineering, State University, 2016'
        ]
    }
}

print("\n" + "="*70)
print("WORD FORMATTING TEST")
print("="*70)

# Get template path
from models.database import TemplateDB
from config import Config

db = TemplateDB()
templates = db.get_all_templates()

if not templates:
    print("❌ No templates found!")
    exit(1)

template = db.get_template(templates[0]['id'])
template_path = os.path.join(Config.TEMPLATE_FOLDER, template['filename'])

print(f"\n✓ Using template: {template['name']}")
print(f"✓ Template path: {template_path}")
print(f"✓ Template exists: {os.path.exists(template_path)}")

# Test opening the template
print(f"\n📄 Opening template...")
doc = Document(template_path)

print(f"✓ Template has {len(doc.paragraphs)} paragraphs")
print(f"✓ Template has {len(doc.tables)} tables")

# Show first 10 paragraphs
print(f"\n📝 First 10 paragraphs:")
for i, para in enumerate(doc.paragraphs[:10]):
    if para.text.strip():
        print(f"  {i}: {para.text[:80]}")

# Check for placeholders
print(f"\n🔍 Looking for placeholders...")
placeholders_found = []

for i, para in enumerate(doc.paragraphs):
    text = para.text
    if '<' in text and '>' in text:
        placeholders_found.append((i, text))
        print(f"  📍 Paragraph {i}: {text[:80]}")

if not placeholders_found:
    print("  ⚠️  No angle bracket placeholders found!")
    print("  💡 Template might use different placeholder format")

# Test replacement
print(f"\n🧪 Testing replacement...")

from utils.word_formatter import format_word_document

template_analysis = {
    'template_path': template_path,
    'template_type': 'docx'
}

output_path = os.path.join(Config.OUTPUT_FOLDER, 'test_output.docx')

print(f"\n🎯 Running formatter...")
success = format_word_document(test_resume_data, template_analysis, output_path)

if success:
    print(f"\n✅ SUCCESS!")
    print(f"📁 Output: {output_path}")
    print(f"\n💡 Open the file to check if placeholders were replaced")
else:
    print(f"\n❌ FAILED!")

print("\n" + "="*70 + "\n")
