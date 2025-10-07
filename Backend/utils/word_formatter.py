"""
Enhanced Word Document Formatter
Handles both .doc and .docx templates
Preserves all formatting, images, headers, footers
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import os
import re
import shutil
import traceback

# Try to import win32com for .doc support
try:
    import win32com.client
    HAS_WIN32 = True
except ImportError:
    HAS_WIN32 = False
    print("⚠️  win32com not available - .doc files will have limited support")

class WordFormatter:
    """Enhanced Word document formatting"""
    
    def __init__(self, resume_data, template_analysis, output_path):
        self.resume_data = resume_data
        self.template_analysis = template_analysis
        self.output_path = output_path
        self.template_path = template_analysis.get('template_path')
        self.template_type = template_analysis.get('template_type')
        
    def format(self):
        """Main formatting method"""
        print(f"\n{'='*70}")
        print(f"📝 WORD DOCUMENT FORMATTING")
        print(f"{'='*70}\n")
        
        print(f"📄 Template: {os.path.basename(self.template_path)}")
        print(f"👤 Candidate: {self.resume_data['name']}")
        print(f"📁 Output: {os.path.basename(self.output_path)}\n")
        
        try:
            # Handle .doc files
            if self.template_path.lower().endswith('.doc'):
                return self._format_doc_file()
            else:
                return self._format_docx_file()
                
        except Exception as e:
            print(f"❌ Error formatting Word document: {e}")
            traceback.print_exc()
            return False
    
    def _format_doc_file(self):
        """Handle .doc files (old Word format)"""
        print("📋 Processing .doc file (old Word format)...")
        
        if HAS_WIN32:
            # Convert .doc to .docx first
            print("✓ Converting .doc to .docx...")
            docx_path = self._convert_doc_to_docx(self.template_path)
            
            if docx_path:
                # Update template path temporarily
                original_path = self.template_path
                self.template_path = docx_path
                
                # Format the docx
                result = self._format_docx_file()
                
                # Cleanup
                try:
                    os.remove(docx_path)
                except:
                    pass
                
                self.template_path = original_path
                return result
            else:
                print("❌ Failed to convert .doc to .docx")
                return False
        else:
            print("⚠️  Cannot process .doc files without win32com")
            print("💡 Please convert template to .docx format or install pywin32")
            return False
    
    def _convert_doc_to_docx(self, doc_path):
        """Convert .doc to .docx using Word COM"""
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            # Open .doc file
            doc = word.Documents.Open(os.path.abspath(doc_path))
            
            # Save as .docx
            docx_path = doc_path.replace('.doc', '_temp.docx')
            doc.SaveAs2(os.path.abspath(docx_path), FileFormat=16)  # 16 = docx format
            
            doc.Close()
            word.Quit()
            
            print(f"✓ Converted to: {docx_path}")
            return docx_path
            
        except Exception as e:
            print(f"❌ Conversion error: {e}")
            return None
    
    def _format_docx_file(self):
        """Format .docx file"""
        print("📋 Processing .docx file...")
        
        # Open template
        doc = Document(self.template_path)
        
        print(f"✓ Template loaded: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
        
        # Show what data we have from resume
        print(f"\n📊 Resume Data Available:")
        print(f"  • Name: {self.resume_data.get('name', 'NOT FOUND')}")
        print(f"  • Email: {self.resume_data.get('email', 'NOT FOUND')}")
        print(f"  • Phone: {self.resume_data.get('phone', 'NOT FOUND')}")
        print(f"  • Experience entries: {len(self.resume_data.get('experience', []))}")
        print(f"  • Education entries: {len(self.resume_data.get('education', []))}")
        print(f"  • Skills: {len(self.resume_data.get('skills', []))}")
        print(f"  • Sections: {list(self.resume_data.get('sections', {}).keys())}")
        
        # Create comprehensive replacement map
        replacements = self._create_replacement_map()
        print(f"\n📝 Created {len(replacements)} replacement mappings")
        
        # Replace in all paragraphs
        replaced_count = 0
        print(f"\n🔍 Scanning {len(doc.paragraphs)} paragraphs for placeholders...")
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            if not paragraph.text.strip():
                continue
                
            # Check each replacement
            for key, value in replacements.items():
                if self._text_contains(paragraph.text, key):
                    print(f"  📍 Found '{key}' in paragraph {para_idx}: '{paragraph.text[:50]}...'")
                    count = self._replace_in_paragraph(paragraph, key, value)
                    if count > 0:
                        print(f"  ✅ Replaced with: '{value[:50]}...'")
                    else:
                        print(f"  ⚠️  Found but couldn't replace (might be in multiple runs)")

            # Regex-driven fallback for angle bracket placeholders with variations
            # Candidate name generic patterns - very flexible to catch all variations
            name_patterns = [
                r"<\s*Candidate[’']?s\s+full\s+name\s*>",
                r"<\s*Candidate\s*Name\s*>",
                r"<\s*Name\s*>",
                r"<\s*[Pp]lease\s+[Ii]nsert\s+[Cc]andidate[’']?s?\s+[Nn]ame\s+[Hh]ere\s*>",
                r"<\s*[Ii]nsert\s+[Cc]andidate[’']?s?\s+[Nn]ame\s+[Hh]ere\s*>",
                r"<\s*[Ii]nsert\s+[Cc]andidate\s+[Nn]ame\s*>",
                r"<\s*[Pp]lease\s+[Ii]nsert\s+[Nn]ame\s*>",
                r"<\s*[Cc]andidate\s+[Nn]ame\s+[Hh]ere\s*>",
            ]
            for pat in name_patterns:
                if re.search(pat, paragraph.text, re.IGNORECASE):
                    before = paragraph.text
                    self._regex_replace_paragraph(paragraph, pat, self.resume_data.get('name', '').strip() or 'Candidate Name')
                    if paragraph.text != before:
                        print(f"  ✅ Regex replaced candidate name in paragraph {para_idx}")
                        replaced_count += 1

            # Generic catch-all: any <...> containing both 'candidate' and 'name' (any order)
            generic_name_pat = r"<[^>]*?(candidate[^>]*name|name[^>]*candidate)[^>]*?>"
            if re.search(generic_name_pat, paragraph.text, re.IGNORECASE):
                before = paragraph.text
                self._regex_replace_paragraph(paragraph, generic_name_pat, self.resume_data.get('name', '').strip() or 'Candidate Name')
                if paragraph.text != before:
                    print(f"  ✅ Generic regex replaced candidate name in paragraph {para_idx}")
                    replaced_count += 1

            # Employment placeholder generic patterns (very flexible)
            emp_patterns = [
                r"<[^>]*employment[^>]*history[^>]*>",
                r"<[^>]*work[^>]*history[^>]*>",
                r"<[^>]*professional[^>]*experience[^>]*>",
                r"<[^>]*career[^>]*(history|experience)[^>]*>",
                r"<[^>]*history[^>]*(employ|employer|work|career)[^>]*>",
                r"<[^>]*list[^>]*employment[^>]*history[^>]*>",
            ]
            for emp_pat in emp_patterns:
                if re.search(emp_pat, paragraph.text, re.IGNORECASE):
                    content = self._find_matching_resume_section('experience', self.resume_data.get('sections', {}))
                    if content:
                        bullets = []
                        for item in content[:10]:
                            if item.strip():
                                bullets.append('• ' + item.strip().lstrip('•').strip())
                        self._regex_replace_paragraph(paragraph, emp_pat, '\n'.join(bullets))
                        print(f"  ✅ Regex replaced experience placeholder in paragraph {para_idx}")
                        replaced_count += 1
                        break

            # Education placeholder generic pattern
            edu_pat = r"<[^>]*education[^>]*background[^>]*>"
            if re.search(edu_pat, paragraph.text, re.IGNORECASE):
                content = self._find_matching_resume_section('education', self.resume_data.get('sections', {}))
                if content:
                    bullets = []
                    for item in content[:10]:
                        if item.strip():
                            bullets.append('• ' + item.strip().lstrip('•').strip())
                    self._regex_replace_paragraph(paragraph, edu_pat, '\n'.join(bullets))
                    print(f"  ✅ Regex replaced education placeholder in paragraph {para_idx}")
        
        print(f"\n✓ Replaced {replaced_count} placeholders in paragraphs")
        
        # Replace in tables and detect skills tables
        table_replaced = 0
        print(f"\n🔍 Scanning {len(doc.tables)} tables...")
        
        for table_idx, table in enumerate(doc.tables):
            # Check if this is a skills table
            if self._is_skills_table(table):
                print(f"  📊 Found skills table at index {table_idx}")
                skills_filled = self._fill_skills_table(table)
                print(f"  ✅ Filled {skills_filled} skill rows")
                table_replaced += skills_filled
            else:
                # Regular placeholder replacement in non-skills tables
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key, value in replacements.items():
                                if self._text_contains(paragraph.text, key):
                                    table_replaced += self._replace_in_paragraph(paragraph, key, value)
        
        print(f"✓ Replaced {table_replaced} placeholders in tables")
        
        # Replace in headers/footers
        header_footer_replaced = 0
        for section in doc.sections:
            # Header
            for paragraph in section.header.paragraphs:
                for key, value in replacements.items():
                    if self._text_contains(paragraph.text, key):
                        header_footer_replaced += self._replace_in_paragraph(paragraph, key, value)
            
            # Footer
            for paragraph in section.footer.paragraphs:
                for key, value in replacements.items():
                    if self._text_contains(paragraph.text, key):
                        header_footer_replaced += self._replace_in_paragraph(paragraph, key, value)
        
        if header_footer_replaced > 0:
            print(f"✓ Replaced {header_footer_replaced} placeholders in headers/footers")
        
        # Add sections content
        sections_added = self._add_sections_content(doc)
        print(f"✓ Added {sections_added} sections")
        
        # Save output
        output_docx = self.output_path.replace('.pdf', '.docx')
        doc.save(output_docx)
        
        print(f"\n✅ Successfully created formatted document!")
        print(f"📁 Saved to: {output_docx}\n")
        
        # Optionally convert to PDF
        if self.output_path.endswith('.pdf'):
            print("📄 Converting to PDF...")
            if self._convert_to_pdf(output_docx, self.output_path):
                print(f"✓ PDF created: {self.output_path}")
                # Keep both docx and pdf
            else:
                print("⚠️  PDF conversion failed, keeping .docx file")
        
        return True

    # Helper: insert a new paragraph directly after a given paragraph
    def _insert_paragraph_after(self, paragraph, text):
        try:
            new_p = OxmlElement('w:p')
            paragraph._p.addnext(new_p)
            new_para = Paragraph(new_p, paragraph._parent)
            new_para.add_run(text)
            return new_para
        except Exception:
            # Fallback: append to document if direct insert fails
            return paragraph._parent.add_paragraph(text)
    
    def _insert_experience_block(self, doc, after_paragraph, exp_data):
        """Insert a structured 2-column experience block"""
        try:
            # Get parsed data from resume parser
            company = exp_data.get('company', '')
            role = exp_data.get('role', '')
            duration = exp_data.get('duration', '')
            details = exp_data.get('details', [])
            
            # Fallback: if company/role not parsed, try to extract from title
            if not company and not role:
                title = exp_data.get('title', '')
                company, role = self._parse_company_role(title)
            
            # Clean up duration format
            duration_clean = self._clean_duration(duration)
            
            # Create 2-column table (no borders) - insert after the heading paragraph
            table = self._insert_table_after(doc, after_paragraph, rows=1, cols=2)
            if not table:
                return False
            
            # Set column widths (70% left, 30% right)
            table.columns[0].width = Inches(4.7)
            table.columns[1].width = Inches(1.8)
            
            # Remove borders
            for row in table.rows:
                for cell in row.cells:
                    self._remove_cell_borders(cell)
            
            # Left cell: COMPANY – ROLE (all uppercase, bold)
            left_cell = table.rows[0].cells[0]
            left_para = left_cell.paragraphs[0]
            
            # Format: COMPANY – ROLE
            if company and role:
                text = f"{company.upper()} – {role.upper()}"
            elif company:
                text = company.upper()
            elif role:
                text = role.upper()
            else:
                text = title.upper()
            
            left_run = left_para.add_run(text)
            left_run.bold = True
            left_run.font.size = Pt(11)
            
            # Right cell: Duration (aligned right)
            right_cell = table.rows[0].cells[1]
            right_para = right_cell.paragraphs[0]
            right_run = right_para.add_run(duration_clean)
            right_run.font.size = Pt(10)
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Add details below if any (in a new row or separate paragraph)
            if details and len(details) > 0:
                # Add a new row for details
                detail_row = table.add_row()
                detail_cell = detail_row.cells[0]
                # Merge cells for full width
                detail_cell.merge(detail_row.cells[1])
                detail_para = detail_cell.paragraphs[0]
                
                # Add details as bullets
                for detail in details[:6]:
                    if detail.strip():
                        detail_text = detail.strip()
                        if not detail_text.startswith('•') and not detail_text.startswith('-'):
                            detail_text = '   • ' + detail_text
                        detail_para.add_run(detail_text + '\n')
            
            return table
            
        except Exception as e:
            print(f"  ⚠️  Error inserting experience block: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def _insert_education_block(self, doc, after_paragraph, edu_data):
        """Insert a structured 2-column education block"""
        try:
            # Get parsed data from resume parser
            degree = edu_data.get('degree', '')
            institution = edu_data.get('institution', '')
            year = edu_data.get('year', '')
            details = edu_data.get('details', [])
            
            # Fallback: if institution not parsed, try to extract from degree or details
            if not institution:
                institution = self._extract_institution(degree, details)
            
            # Clean up year format
            year_clean = self._clean_duration(year)
            
            # Create 2-column table (no borders) - insert after the heading paragraph
            table = self._insert_table_after(doc, after_paragraph, rows=1, cols=2)
            if not table:
                return None
            
            # Set column widths (70% left, 30% right)
            table.columns[0].width = Inches(4.7)
            table.columns[1].width = Inches(1.8)
            
            # Remove borders
            for row in table.rows:
                for cell in row.cells:
                    self._remove_cell_borders(cell)
            
            # Left cell: DEGREE  INSTITUTION (uppercase, bold)
            left_cell = table.rows[0].cells[0]
            left_para = left_cell.paragraphs[0]
            
            # Format: DEGREE  INSTITUTION
            if degree and institution:
                text = f"{degree.upper()}  {institution.upper()}"
            elif degree:
                text = degree.upper()
            elif institution:
                text = institution.upper()
            else:
                text = "EDUCATION"
            
            left_run = left_para.add_run(text)
            left_run.bold = True
            left_run.font.size = Pt(11)
            
            # Right cell: Year (aligned right)
            right_cell = table.rows[0].cells[1]
            right_para = right_cell.paragraphs[0]
            right_run = right_para.add_run(year_clean)
            right_run.font.size = Pt(10)
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Add details below if any (in a new row)
            if details and len(details) > 0:
                # Add a new row for details
                detail_row = table.add_row()
                detail_cell = detail_row.cells[0]
                # Merge cells for full width
                detail_cell.merge(detail_row.cells[1])
                detail_para = detail_cell.paragraphs[0]
                
                # Add details
                for detail in details[:3]:
                    if detail.strip() and detail.strip().lower() != institution.lower():
                        detail_para.add_run('   - ' + detail.strip() + '\n')
            
            return table
            
        except Exception as e:
            print(f"  ⚠️  Error inserting education block: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def _clean_duration(self, duration):
        """Clean and format duration string to YYYY-YYYY (or single YYYY)"""
        if not duration:
            return ''

        duration = duration.strip()

        # Normalize common separators to '-'
        duration = re.sub(r'\s*(to|–|—|-)\s*', '-', duration, flags=re.IGNORECASE)

        # Replace words with years
        current_year = '2025'
        duration = re.sub(r'\b(current|present)\b', current_year, duration, flags=re.IGNORECASE)

        # Map months to years where patterns like 'Apr 2013' appear
        # Extract full 4-digit years
        years = re.findall(r'\b(?:19|20)\d{2}\b', duration)

        if len(years) >= 2:
            return f"{years[0]}-{years[-1]}"
        elif len(years) == 1:
            return years[0]
        else:
            return ''
    
    def _parse_company_role(self, title):
        """Parse company and role from title line"""
        # Common patterns: "Company Name - Role" or "Role at Company" or "Role, Company"
        if ' - ' in title:
            parts = title.split(' - ', 1)
            return parts[0].strip(), parts[1].strip()
        elif ' at ' in title.lower():
            parts = re.split(r'\s+at\s+', title, flags=re.IGNORECASE)
            return parts[1].strip() if len(parts) > 1 else '', parts[0].strip()
        elif ', ' in title:
            parts = title.split(', ', 1)
            return parts[1].strip(), parts[0].strip()
        else:
            # Assume entire line is company or role
            return title.strip(), ''
    
    def _extract_institution(self, degree, details):
        """Extract institution name from degree line or details"""
        # Check if degree line contains institution (common pattern: "Degree, Institution")
        if ', ' in degree:
            parts = degree.split(', ', 1)
            return parts[1].strip()
        
        # Look in details for institution keywords
        institution_keywords = ['university', 'college', 'institute', 'school', 'academy']
        for detail in details:
            if any(kw in detail.lower() for kw in institution_keywords):
                return detail.strip()
        
        return ''
    
    def _remove_cell_borders(self, cell):
        """Remove all borders from a table cell"""
        try:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            tcPr.append(tcBorders)
        except:
            pass
    
    def _insert_table_after(self, doc, anchor, rows=1, cols=2):
        """Create a table and position it immediately after the given anchor (Paragraph or Table)."""
        try:
            # Create table using python-docx API
            table = doc.add_table(rows=rows, cols=cols)
            tbl = table._element

            # Resolve anchor element (paragraph or table)
            if hasattr(anchor, '_element'):
                anchor_elm = anchor._element
            else:
                anchor_elm = anchor

            # Move table right after the anchor in the document body
            anchor_elm.addnext(tbl)

            return table

        except Exception as e:
            print(f"  ⚠️  Error creating table: {e}")
            import traceback
            traceback.print_exc()
            return None

    def _delete_following_bullets(self, paragraph, max_scan=40):
        """Delete bullet-like paragraphs after a heading/placeholder to avoid duplicate raw lists.
        Only removes paragraphs that look like bullets or numbered lists; stops at next heading/table.
        """
        try:
            body = paragraph._element.getparent()
            node = paragraph._element.getnext()
            scanned = 0
            while node is not None and scanned < max_scan:
                scanned += 1
                if node.tag.endswith('tbl'):
                    break
                if node.tag.endswith('p'):
                    # Extract plain text
                    text_nodes = node.xpath('.//w:t', namespaces=node.nsmap) if hasattr(node, 'xpath') else []
                    text = ''.join([t.text for t in text_nodes if t is not None and t.text is not None])
                    txt = (text or '').strip()
                    norm = txt.upper()
                    # Stop at next heading keywords
                    if any(k in norm for k in ['EDUCATION', 'SKILLS', 'SUMMARY', 'PROJECT', 'CERTIFICATION', 'EXPERIENCE', 'WORK EXPERIENCE', 'EMPLOYMENT HISTORY']):
                        break
                    # Bullet/numbered heuristics
                    is_bullet = txt.startswith(('•', '-', '–', '—', '*', '●'))
                    is_numbered = bool(re.match(r'^\d+[\).\-\s]', txt))
                    if is_bullet or is_numbered:
                        nxt = node.getnext()
                        body.remove(node)
                        node = nxt
                        continue
                    else:
                        # If we hit a non-bullet normal paragraph, stop cleanup
                        break
                node = node.getnext()
        except Exception:
            pass
    
    def _collect_bullets_after_heading(self, paragraph, max_scan=50):
        """Collect consecutive bullet-like paragraphs immediately after a heading/placeholder."""
        bullets = []
        try:
            node = paragraph._element.getnext()
            scanned = 0
            while node is not None and scanned < max_scan:
                scanned += 1
                if node.tag.endswith('tbl'):
                    # Collect all paragraph texts from the table as bullets
                    paras = node.xpath('.//w:p', namespaces=node.nsmap) if hasattr(node, 'xpath') else []
                    for p in paras:
                        tnodes = p.xpath('.//w:t', namespaces=p.nsmap) if hasattr(p, 'xpath') else []
                        text = ''.join([t.text for t in tnodes if t is not None and t.text is not None]).strip()
                        if text:
                            bullets.append(text.lstrip(' •–—-*●'))
                    break
                if node.tag.endswith('p'):
                    text_nodes = node.xpath('.//w:t', namespaces=node.nsmap) if hasattr(node, 'xpath') else []
                    text = ''.join([t.text for t in text_nodes if t is not None and t.text is not None])
                    txt = (text or '').strip()
                    norm = txt.upper()
                    if any(k in norm for k in ['EDUCATION', 'SKILLS', 'SUMMARY', 'PROJECT', 'CERTIFICATION', 'EXPERIENCE', 'WORK EXPERIENCE', 'EMPLOYMENT HISTORY']):
                        break
                    if txt.startswith(('•', '-', '–', '—', '*', '●')) or re.match(r'^\d+[\).\-\s]', txt):
                        bullets.append(txt.lstrip(' •–—-*●'))
                    else:
                        break
                node = node.getnext()
        except Exception:
            pass
        return bullets

    def _delete_next_table(self, paragraph):
        """Delete the immediate next table after a heading/placeholder (used when raw content is a table)."""
        try:
            node = paragraph._element.getnext()
            if node is not None and node.tag.endswith('tbl'):
                parent = node.getparent()
                parent.remove(node)
        except Exception:
            pass
    
    def _build_experience_from_bullets(self, bullets):
        """Best-effort convert raw bullet lines into structured exp list when parser is empty."""
        exps = []
        i = 0
        while i < len(bullets):
            line = bullets[i]
            # Case: role + dates on this line
            if re.search(r'(?:19|20)\d{2}', line):
                duration = self._clean_duration(line)
                role = re.sub(r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-zA-Z]*\s+(?:19|20)\d{2}\b', '', line, flags=re.IGNORECASE)
                role = re.sub(r'\b(?:19|20)\d{2}\b', '', role)
                role = re.sub(r'\b(to|–|—|-)\b', '', role, flags=re.IGNORECASE).strip(' ,;:-')
                # Next bullet as company if available
                company = ''
                if i + 1 < len(bullets):
                    company = bullets[i+1]
                    # Strip obvious location fragments
                    company = re.sub(r',[^,]*\b(?:city|state|india|usa|uk)\b.*$', '', company, flags=re.IGNORECASE).strip()
                details = []
                j = i + 2
                while j < len(bullets):
                    if re.search(r'(?:19|20)\d{2}', bullets[j]):
                        break
                    details.append(bullets[j])
                    j += 1
                exps.append({'company': company, 'role': role, 'duration': duration, 'details': details})
                i = j
            else:
                i += 1
        return exps
    
    def _build_education_from_bullets(self, bullets):
        """Convert raw education bullets into degree/institution/year list when parser is empty."""
        edus = []
        i = 0
        while i < len(bullets):
            line = bullets[i]
            degree = ''
            institution = ''
            year = ''
            # Try to split by institution keyword
            m = re.search(r'(university|college|school|institute|academy)\b.*', line, flags=re.IGNORECASE)
            if m:
                degree = line[:m.start()].strip(' ,;:-')
                institution = line[m.start():].strip()
                year = self._clean_duration(line)
            else:
                # If next line is year, treat current as degree+institution
                if i + 1 < len(bullets) and re.search(r'(?:19|20)\d{2}', bullets[i+1]):
                    degree = line
                    year = self._clean_duration(bullets[i+1])
                    i += 1
                else:
                    degree = line
                    year = self._clean_duration(line)
            # Cleanup
            degree = degree.strip()
            institution = re.sub(r',[^,]*\b(?:city|state|india|usa|uk)\b.*$', '', institution, flags=re.IGNORECASE).strip()
            edus.append({'degree': degree, 'institution': institution, 'year': year, 'details': []})
            i += 1
        return edus
    
    def _create_replacement_map(self):
        """Create comprehensive replacement map"""
        replacements = {}
        
        # Personal information - Multiple formats
        # NOTE: Be specific to avoid replacing CAI contact manager info
        if self.resume_data.get('name'):
            replacements['[NAME]'] = self.resume_data['name']
            replacements['[CANDIDATE NAME]'] = self.resume_data['name']
            replacements['<CANDIDATE NAME>'] = self.resume_data['name']
            replacements["<Candidate's full name>"] = self.resume_data['name']
            replacements['<Candidate Name>'] = self.resume_data['name']
            replacements['<Name>'] = self.resume_data['name']
            replacements['Your Name'] = self.resume_data['name']
            replacements['CANDIDATE NAME'] = self.resume_data['name']
            # DO NOT replace "Insert name" as it might be in CAI contact section
        
        if self.resume_data.get('email'):
            # ONLY replace explicit placeholders, NOT actual email addresses
            replacements['[EMAIL]'] = self.resume_data['email']
            replacements['[Email]'] = self.resume_data['email']
            replacements['<EMAIL>'] = self.resume_data['email']
            replacements['<Email>'] = self.resume_data['email']
            replacements['<Candidate Email>'] = self.resume_data['email']
            # DO NOT replace example emails or "Email:" labels to avoid changing CAI contact info
        
        if self.resume_data.get('phone'):
            # ONLY replace explicit placeholders, NOT actual phone numbers
            replacements['[PHONE]'] = self.resume_data['phone']
            replacements['[Phone]'] = self.resume_data['phone']
            replacements['<PHONE>'] = self.resume_data['phone']
            replacements['<Phone>'] = self.resume_data['phone']
            replacements['<Candidate Phone>'] = self.resume_data['phone']
            # DO NOT replace example numbers or "Phone:" labels to avoid changing CAI contact info
        
        if self.resume_data.get('address'):
            replacements['[ADDRESS]'] = self.resume_data['address']
            replacements['[Address]'] = self.resume_data['address']
            replacements['<ADDRESS>'] = self.resume_data['address']
            replacements['<Address>'] = self.resume_data['address']
            replacements['Your Address'] = self.resume_data['address']
        
        if self.resume_data.get('linkedin'):
            replacements['[LINKEDIN]'] = self.resume_data['linkedin']
            replacements['[LinkedIn]'] = self.resume_data['linkedin']
            replacements['<LINKEDIN>'] = self.resume_data['linkedin']
            replacements['<LinkedIn>'] = self.resume_data['linkedin']
            replacements['linkedin.com/in/username'] = self.resume_data['linkedin']
        
        if self.resume_data.get('dob'):
            replacements['[DOB]'] = self.resume_data['dob']
            replacements['[Date of Birth]'] = self.resume_data['dob']
            replacements['<DOB>'] = self.resume_data['dob']
        
        return replacements
    
    def _text_contains(self, text, search_term):
        """Case-insensitive text search"""
        return search_term.lower() in text.lower()
    
    def _replace_in_paragraph(self, paragraph, search_term, replacement):
        """Replace text in paragraph while preserving formatting"""
        replaced = 0
        
        # First try: Replace in individual runs
        for run in paragraph.runs:
            if self._text_contains(run.text, search_term):
                # Case-insensitive replacement
                pattern = re.compile(re.escape(search_term), re.IGNORECASE)
                run.text = pattern.sub(replacement, run.text)
                replaced += 1
        
        # Second try: If not found in individual runs, text might be split
        # Combine all runs and check
        if replaced == 0 and self._text_contains(paragraph.text, search_term):
            # Text is split across runs - need to handle differently
            full_text = paragraph.text
            pattern = re.compile(re.escape(search_term), re.IGNORECASE)
            new_text = pattern.sub(replacement, full_text)
            
            if new_text != full_text:
                # Clear all runs and add new text
                for run in paragraph.runs:
                    run.text = ''
                
                # Add replacement text to first run
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                    replaced += 1
                else:
                    # No runs, add new run
                    paragraph.add_run(new_text)
                    replaced += 1
        
        return replaced

    def _regex_replace_paragraph(self, paragraph, pattern, replacement):
        """Regex-based replacement across runs: rebuilds paragraph text."""
        try:
            full_text = paragraph.text or ''
            new_text = re.sub(pattern, replacement, full_text, flags=re.IGNORECASE)
            if new_text != full_text:
                # clear runs and set new_text
                for run in paragraph.runs:
                    run.text = ''
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                else:
                    paragraph.add_run(new_text)
        except Exception:
            pass
    
    def _add_sections_content(self, doc):
        """Add resume sections to document and replace placeholders"""
        sections_added = 0
        
        # Section placeholder patterns
        section_placeholders = {
            'experience': [
                "<List candidate's relevant employment history>",
                '<List employment history>',
                '<Employment History>',
                '<Work Experience>',
                '<Professional Experience>',
                '<Career History>',
                '<History of Employment>',
                '<History of the Employer>',
                '<Employer History>',
                '<Experience details>',
                'List relevant employment history',
                'List the employment history'
            ],
            'education': [
                "<List candidate's education background>",
                '<List education background>',
                '<Education Background>',
                '<Education details>',
                'List education background'
            ],
            'skills': [
                '<List skills>',
                '<Skills>',
                '<Technical Skills>',
                'List skills'
            ],
            'summary': [
                '<Professional Summary>',
                '<Summary>',
                '<Objective>',
                'Professional summary'
            ]
        }
        
        resume_sections = self.resume_data.get('sections', {})
        
        # Replace section placeholders with actual content
        print(f"\n🔍 Looking for section placeholders...")
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text.strip()
            
            # Check for section placeholders
            for section_key, placeholders in section_placeholders.items():
                for placeholder in placeholders:
                    if placeholder.lower() in para_text.lower():
                        print(f"  📍 Found placeholder: '{placeholder}' in paragraph {para_idx}")
                        
                        if section_key == 'experience':
                            experiences = self.resume_data.get('experience', [])
                            if experiences:
                                # Clear placeholder paragraph
                                for run in paragraph.runs:
                                    run.text = ''
                                # Remove any existing bullet lines under this heading/placeholder
                                self._delete_following_bullets(paragraph)
                                # Insert structured blocks after this paragraph
                                last_element = paragraph
                                for exp in experiences[:10]:
                                    table = self._insert_experience_block(doc, last_element, exp)
                                    if table:
                                        last_element = table.rows[0].cells[0].paragraphs[0]
                                sections_added += 1
                                print(f"  ✅ Inserted {len(experiences[:10])} structured experience block(s)")
                                continue
                        
                        if section_key == 'education':
                            education = self.resume_data.get('education', [])
                            if education:
                                # Clear placeholder paragraph
                                for run in paragraph.runs:
                                    run.text = ''
                                # Remove any existing bullet lines under this heading/placeholder
                                self._delete_following_bullets(paragraph)
                                # Insert structured blocks after this paragraph
                                last_element = paragraph
                                for edu in education[:5]:
                                    table = self._insert_education_block(doc, last_element, edu)
                                    if table:
                                        last_element = table.rows[0].cells[0].paragraphs[0]
                                sections_added += 1
                                print(f"  ✅ Inserted {len(education[:5])} structured education block(s)")
                                continue
                        
                        # Fallback for other sections: bullets
                        content = self._find_matching_resume_section(section_key, resume_sections)
                        if content and len(content) > 0:
                            content_lines = []
                            for item in content[:10]:  # Limit to 10 items
                                if item.strip():
                                    if not item.strip().startswith('•'):
                                        content_lines.append(f'• {item.strip()}')
                                    else:
                                        content_lines.append(item.strip())
                            content_text = '\n'.join(content_lines)
                            full_text = paragraph.text
                            pattern = re.compile(re.escape(placeholder), re.IGNORECASE)
                            new_text = pattern.sub(content_text, full_text)
                            if new_text != full_text:
                                for run in paragraph.runs:
                                    run.text = ''
                                if paragraph.runs:
                                    paragraph.runs[0].text = new_text
                                else:
                                    paragraph.add_run(new_text)
                                sections_added += 1
                                print(f"  ✅ Replaced with {len(content_lines)} items from resume")
                        else:
                            print(f"  ⚠️  No matching content found in resume for '{section_key}'")
        
        # Also look for section headings and add content after them
        section_markers = {
            'experience': ['EMPLOYMENT HISTORY', 'WORK EXPERIENCE', 'PROFESSIONAL EXPERIENCE', 'EXPERIENCE', 'CAREER HISTORY', 'HISTORY OF EMPLOYMENT', 'HISTORY OF THE EMPLOYER', 'EMPLOYER HISTORY'],
            'education': ['EDUCATION', 'ACADEMIC BACKGROUND', 'QUALIFICATIONS'],
            'skills': ['SKILLS', 'TECHNICAL SKILLS', 'CORE COMPETENCIES'],
            'summary': ['SUMMARY', 'PROFESSIONAL SUMMARY', 'OBJECTIVE', 'PROFILE'],
        }
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            para_text_upper = paragraph.text.upper().strip()
            
            for section_key, markers in section_markers.items():
                # Match either exact heading or heading containing the marker text
                if any((marker == para_text_upper) or (marker in para_text_upper) for marker in markers):
                    # Found a section heading
                    
                    # For experience and education, use structured blocks
                    if section_key == 'experience':
                        experiences = self.resume_data.get('experience', [])
                        # Fallback: derive from raw bullets if parser returned none
                        if not experiences:
                            raw_bullets = self._collect_bullets_after_heading(paragraph)
                            experiences = self._build_experience_from_bullets(raw_bullets)
                        if experiences:
                            print(f"  • Found EXPERIENCE heading: {paragraph.text} → inserting {len(experiences)} structured block(s)")
                            # Remove any existing bullet lines under this heading before inserting
                            self._delete_following_bullets(paragraph)
                            last_element = paragraph
                            for exp in experiences[:10]:  # Limit to 10 experiences
                                table = self._insert_experience_block(doc, last_element, exp)
                                if table:
                                    # Update last_element to be the table so next one inserts after it
                                    last_element = table.rows[0].cells[0].paragraphs[0]
                            sections_added += 1
                    
                    elif section_key == 'education':
                        education = self.resume_data.get('education', [])
                        # Fallback: derive from raw bullets if parser returned none
                        if not education:
                            raw_bullets = self._collect_bullets_after_heading(paragraph)
                            education = self._build_education_from_bullets(raw_bullets)
                        if education:
                            print(f"  • Found EDUCATION heading: {paragraph.text} → inserting {len(education)} structured block(s)")
                            # Remove any existing bullet lines under this heading before inserting
                            self._delete_following_bullets(paragraph)
                            last_element = paragraph
                            for edu in education[:5]:  # Limit to 5 education entries
                                table = self._insert_education_block(doc, last_element, edu)
                                if table:
                                    # Update last_element to be the table so next one inserts after it
                                    last_element = table.rows[0].cells[0].paragraphs[0]
                            sections_added += 1
                    
                    else:
                        # For other sections (skills, summary), use simple bullets
                        content = self._find_matching_resume_section(section_key, resume_sections)
                        if content:
                            print(f"  • Found section heading: {paragraph.text} → inserting {len(content[:10])} item(s)")
                            insert_after = paragraph
                            for item in content[:10]:
                                txt = item.strip()
                                if not txt:
                                    continue
                                if not txt.startswith('•'):
                                    txt = f"• {txt}"
                                insert_after = self._insert_paragraph_after(insert_after, txt)
                            sections_added += 1
            
            # Generic heading catch-all: lines containing both HISTORY and (EMPLOY/EMPLOYER/WORK/CAREER)
            if 'HISTORY' in para_text_upper and any(k in para_text_upper for k in ['EMPLOY', 'EMPLOYER', 'WORK', 'CAREER']):
                experiences = self.resume_data.get('experience', [])
                if experiences:
                    print(f"  • Found generic experience heading: {paragraph.text} → inserting {len(experiences)} structured block(s)")
                    # Remove any existing bullet lines under this heading before inserting
                    self._delete_following_bullets(paragraph)
                    last_element = paragraph
                    for exp in experiences[:10]:
                        table = self._insert_experience_block(doc, last_element, exp)
                        if table:
                            last_element = table.rows[0].cells[0].paragraphs[0]
                    sections_added += 1
        
        # PASS 3: Scan tables for headings inside cells and insert structured blocks there too
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        para_text_upper = (paragraph.text or '').upper().strip()
                        # EXPERIENCE in tables
                        if any(m in para_text_upper for m in ['EMPLOYMENT HISTORY', 'WORK EXPERIENCE', 'PROFESSIONAL EXPERIENCE', 'EXPERIENCE', 'CAREER HISTORY', 'HISTORY OF EMPLOYMENT', 'HISTORY OF THE EMPLOYER', 'EMPLOYER HISTORY']):
                            experiences = self.resume_data.get('experience', [])
                            if not experiences:
                                raw_bullets = self._collect_bullets_after_heading(paragraph)
                                experiences = self._build_experience_from_bullets(raw_bullets)
                            if experiences:
                                self._delete_following_bullets(paragraph)
                                self._delete_next_table(paragraph)
                                last_element = paragraph
                                for exp in experiences[:10]:
                                    tbl = self._insert_experience_block(doc, last_element, exp)
                                    if tbl:
                                        last_element = tbl.rows[0].cells[0].paragraphs[0]
                                sections_added += 1
                        # EDUCATION in tables
                        if any(m in para_text_upper for m in ['EDUCATION', 'ACADEMIC BACKGROUND', 'QUALIFICATIONS']):
                            education = self.resume_data.get('education', [])
                            if not education:
                                raw_bullets = self._collect_bullets_after_heading(paragraph)
                                education = self._build_education_from_bullets(raw_bullets)
                            if education:
                                self._delete_following_bullets(paragraph)
                                self._delete_next_table(paragraph)
                                last_element = paragraph
                                for edu in education[:5]:
                                    tbl = self._insert_education_block(doc, last_element, edu)
                                    if tbl:
                                        last_element = tbl.rows[0].cells[0].paragraphs[0]
                                sections_added += 1
        
        return sections_added
    
    def _is_skills_table(self, table):
        """Check if table is a skills table by examining headers"""
        if len(table.rows) < 2:  # Need at least header + 1 row
            return False
        
        # Get first row (header) text
        header_row = table.rows[0]
        header_texts = [cell.text.strip().lower() for cell in header_row.cells]
        
        # Check for skills table indicators
        skills_keywords = ['skill', 'skills', 'technology', 'competency']
        years_keywords = ['years', 'experience', 'years used', 'years of experience']
        last_used_keywords = ['last used', 'last', 'recent', 'most recent']
        
        has_skill_col = any(any(kw in h for kw in skills_keywords) for h in header_texts)
        has_years_col = any(any(kw in h for kw in years_keywords) for h in header_texts)
        has_last_used_col = any(any(kw in h for kw in last_used_keywords) for h in header_texts)
        
        # It's a skills table if it has skill column and at least one other column
        return has_skill_col and (has_years_col or has_last_used_col)
    
    def _fill_skills_table(self, table):
        """Fill skills table with candidate's skills data"""
        if len(table.rows) < 2:
            return 0
        
        # Get header row to identify columns
        header_row = table.rows[0]
        header_texts = [cell.text.strip().lower() for cell in header_row.cells]
        
        # Find column indices
        skill_col = None
        years_col = None
        last_used_col = None
        
        for idx, header in enumerate(header_texts):
            if 'skill' in header or 'technology' in header:
                skill_col = idx
            elif 'years' in header or 'experience' in header:
                years_col = idx
            elif 'last' in header or 'recent' in header:
                last_used_col = idx
        
        if skill_col is None:
            return 0
        
        # Get skills from resume
        skills_data = self._extract_skills_with_details()
        
        if not skills_data:
            return 0
        
        # Clear existing data rows (keep header)
        rows_to_delete = []
        for i in range(1, len(table.rows)):
            rows_to_delete.append(i)
        
        # Delete from bottom to top to avoid index issues
        for i in reversed(rows_to_delete):
            table._element.remove(table.rows[i]._element)
        
        # Add skills rows
        filled_count = 0
        for skill_info in skills_data[:15]:  # Limit to 15 skills
            # Add new row
            new_row = table.add_row()
            
            # Fill skill name
            if skill_col is not None:
                new_row.cells[skill_col].text = skill_info.get('skill', '')
            
            # Fill years
            if years_col is not None:
                new_row.cells[years_col].text = skill_info.get('years', '')
            
            # Fill last used
            if last_used_col is not None:
                new_row.cells[last_used_col].text = skill_info.get('last_used', '')
            
            filled_count += 1
        
        return filled_count
    
    def _extract_skills_with_details(self):
        """Extract skills with years and last used info from resume data"""
        skills_list = []
        
        # Get skills from resume data
        skills = self.resume_data.get('skills', [])
        experience = self.resume_data.get('experience', [])
        
        # Try to extract years from experience
        current_year = 2025
        
        for skill in skills[:15]:  # Limit to 15 skills
            skill_name = skill if isinstance(skill, str) else skill.get('name', '')
            
            # Try to find this skill in experience to get dates
            years_exp = ''
            last_used = ''
            
            # Search through experience for this skill
            for exp in experience:
                exp_text = str(exp).lower()
                if skill_name.lower() in exp_text:
                    # Try to extract years
                    duration = exp.get('duration', '') if isinstance(exp, dict) else ''
                    
                    # Parse years from duration like "2020-2023" or "2020-Present"
                    import re
                    year_matches = re.findall(r'(20\d{2})', str(duration))
                    if year_matches:
                        start_year = int(year_matches[0])
                        end_year = int(year_matches[-1]) if len(year_matches) > 1 else current_year
                        
                        if 'present' in str(duration).lower() or 'current' in str(duration).lower():
                            end_year = current_year
                        
                        years_count = end_year - start_year
                        if years_count > 0:
                            years_exp = f"{years_count}+ years"
                            last_used = str(end_year) if end_year < current_year else "Present"
                    
                    break
            
            # Default values if not found in experience
            if not years_exp:
                years_exp = "1+ years"
            if not last_used:
                last_used = "Recent"
            
            skills_list.append({
                'skill': skill_name,
                'years': years_exp,
                'last_used': last_used
            })
        
        return skills_list
    
    def _find_matching_resume_section(self, section_key, resume_sections):
        """Find matching resume section with synonyms"""
        # Direct match
        if section_key in resume_sections:
            return resume_sections[section_key]

        synonyms = {
            'experience': ['experience', 'employment', 'work', 'professional'],
            'education': ['education', 'academic', 'qualification', 'academics'],
            'skills': ['skills', 'technical', 'competencies', 'expertise'],
            'summary': ['summary', 'objective', 'profile', 'about'],
            'projects': ['projects', 'portfolio'],
            'certifications': ['certifications', 'certificates', 'licenses'],
            'awards': ['awards', 'achievements', 'honors']
        }

        patterns = synonyms.get(section_key, [section_key])
        for resume_key, content in resume_sections.items():
            key_lower = resume_key.lower()
            if any(p in key_lower for p in patterns):
                return content

        return []
    
    def _convert_to_pdf(self, docx_path, pdf_path):
        """Convert DOCX to PDF"""
        try:
            if HAS_WIN32:
                # Use Word COM to convert
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                
                doc = word.Documents.Open(os.path.abspath(docx_path))
                doc.SaveAs2(os.path.abspath(pdf_path), FileFormat=17)  # 17 = PDF format
                doc.Close()
                word.Quit()
                
                return True
            else:
                print("⚠️  PDF conversion requires Microsoft Word")
                return False
                
        except Exception as e:
            print(f"❌ PDF conversion error: {e}")
            return False


def format_word_document(resume_data, template_analysis, output_path):
    """Main function for Word document formatting"""
    formatter = WordFormatter(resume_data, template_analysis, output_path)
    return formatter.format()
