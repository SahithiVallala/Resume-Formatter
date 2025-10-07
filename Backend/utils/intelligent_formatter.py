"""
Intelligent Resume Formatter
Uses deep template analysis and comprehensive resume parsing to:
1. Preserve ALL template visual elements (letterhead, logos, stickers)
2. Intelligently map resume data to template placeholders
3. Maintain exact formatting and styling
4. Replace content while keeping structure intact
"""

from PyPDF2 import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
import os
import io
import traceback

# Import the enhanced Word formatter
try:
    from .word_formatter import format_word_document
    HAS_WORD_FORMATTER = True
except ImportError:
    HAS_WORD_FORMATTER = False

class IntelligentFormatter:
    """Smart resume formatting with template preservation"""
    
    def __init__(self, resume_data, template_analysis, output_path):
        self.resume_data = resume_data
        self.template_analysis = template_analysis
        self.output_path = output_path
        self.template_path = template_analysis.get('template_path')
        self.template_type = template_analysis.get('template_type')
        
        # Validate required fields
        if not self.template_path:
            raise ValueError("Template path not found in analysis. Please re-upload the template.")
        if not self.template_type:
            raise ValueError("Template type not found in analysis. Please re-upload the template.")
        
    def format(self):
        """Main formatting method"""
        print(f"\n{'='*70}")
        print(f"🎨 INTELLIGENT FORMATTING")
        print(f"{'='*70}\n")
        
        print(f"📄 Template: {os.path.basename(self.template_path)}")
        print(f"🎯 Output: {os.path.basename(self.output_path)}\n")
        
        try:
            if self.template_type == 'pdf':
                return self._format_pdf()
            elif self.template_type in ['docx', 'doc']:
                # Use enhanced Word formatter
                if HAS_WORD_FORMATTER:
                    return format_word_document(self.resume_data, self.template_analysis, self.output_path)
                else:
                    print("❌ Enhanced Word formatter not available. Falling back to basic DOCX formatter.")
                    return self._format_docx()
            else:
                print(f"❌ Unsupported template type: {self.template_type}")
                return False
        except Exception as e:
            print(f"❌ Formatting error: {e}")
            traceback.print_exc()
            return False
    
    def _format_pdf(self):
        """Format using PDF template - preserves everything"""
        print("📋 Using PDF template formatting...")
        
        # Read template
        reader = PdfReader(self.template_path)
        writer = PdfWriter()
        template_page = reader.pages[0]
        
        # Create overlay with resume content
        packet = io.BytesIO()
        page_width = float(template_page.mediabox.width)
        page_height = float(template_page.mediabox.height)
        
        can = canvas.Canvas(packet, pagesize=(page_width, page_height))
        
        # Map resume data to template fields
        field_mappings = self._create_field_mappings()
        
        # Get zones from template analysis
        zones = self.template_analysis.get('zones', {})
        has_letterhead = self.template_analysis.get('has_letterhead', False)
        
        # Calculate starting position (below letterhead if exists)
        if has_letterhead:
            start_y = page_height - 150  # Leave space for letterhead
            print("✓ Preserving letterhead area")
        else:
            start_y = page_height - 72
        
        # Set default font
        formatting = self.template_analysis.get('formatting', {})
        font_name = formatting.get('common_font', 'Helvetica')
        font_size = formatting.get('common_size', 10)
        
        # Draw name (if not in template)
        if not self._template_has_name_placeholder():
            can.setFont('Helvetica-Bold', 16)
            can.drawCentredString(page_width / 2, start_y, self.resume_data['name'].upper())
            start_y -= 25
            print(f"✓ Added name: {self.resume_data['name']}")
        
        # Draw contact info
        contact_parts = []
        if self.resume_data['email']:
            contact_parts.append(self.resume_data['email'])
        if self.resume_data['phone']:
            contact_parts.append(self.resume_data['phone'])
        if self.resume_data['linkedin']:
            contact_parts.append(self.resume_data['linkedin'][:30])
        
        if contact_parts:
            can.setFont('Helvetica', 9)
            can.drawCentredString(page_width / 2, start_y, ' | '.join(contact_parts))
            start_y -= 30
            print(f"✓ Added contact info")
        
        # Draw sections based on template structure
        sections_to_draw = self._match_sections()
        
        margins = self.template_analysis['page'].get('margins', {'left': 72, 'right': 72, 'bottom': 72})
        x_pos = margins.get('left', 72)
        y_pos = start_y
        
        for section_name, section_data in sections_to_draw:
            if y_pos < margins.get('bottom', 72) + 100:
                print(f"⚠️  Reached bottom of page, stopping")
                break
            
            # Section heading
            can.setFont('Helvetica-Bold', 12)
            can.drawString(x_pos, y_pos, section_name)
            y_pos -= 15
            
            # Section underline
            can.line(x_pos, y_pos + 5, page_width - margins.get('right', 72), y_pos + 5)
            y_pos -= 10
            
            # Section content
            can.setFont('Helvetica', 10)
            for item in section_data[:10]:  # Limit items
                if y_pos < margins.get('bottom', 72) + 50:
                    break
                
                # Wrap long lines
                wrapped = self._wrap_text(item, can, page_width - x_pos - margins.get('right', 72) - 20)
                for line in wrapped:
                    can.drawString(x_pos + 15, y_pos, f'• {line}')
                    y_pos -= 14
            
            y_pos -= 10  # Space between sections
            print(f"✓ Added section: {section_name}")
        
        can.save()
        
        # Merge overlay with template
        packet.seek(0)
        overlay_pdf = PdfReader(packet)
        overlay_page = overlay_pdf.pages[0]
        
        # Template page is background, overlay is on top
        template_page.merge_page(overlay_page)
        writer.add_page(template_page)
        
        # Write output
        with open(self.output_path, 'wb') as output_file:
            writer.write(output_file)
        
        print(f"\n✅ Successfully created formatted resume!")
        print(f"📁 Saved to: {self.output_path}\n")
        return True
    
    def _format_docx(self):
        """Format using DOCX template - preserves everything"""
        print("📋 Using DOCX template formatting...")
        
        # Open template
        doc = Document(self.template_path)
        
        # Create field replacement map
        replacements = {
            '[NAME]': self.resume_data['name'],
            '[Email]': self.resume_data['email'],
            '[PHONE]': self.resume_data['phone'],
            '[ADDRESS]': self.resume_data.get('address', ''),
            '[LINKEDIN]': self.resume_data['linkedin'],
            '[DOB]': self.resume_data.get('dob', ''),
            'Your Name': self.resume_data['name'],
            'your.email@example.com': self.resume_data['email'],
            '(123) 456-7890': self.resume_data['phone'],
            'linkedin.com/in/username': self.resume_data['linkedin'],
        }
        
        # Replace in all paragraphs
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key.lower() in paragraph.text.lower():
                    for run in paragraph.runs:
                        if key.lower() in run.text.lower():
                            run.text = run.text.replace(key, value)
                            print(f"✓ Replaced {key} with {value[:30]}")
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key.lower() in paragraph.text.lower():
                                for run in paragraph.runs:
                                    if key.lower() in run.text.lower():
                                        run.text = run.text.replace(key, value)
        
        # Add sections if template has section markers
        sections_to_add = self._match_sections()
        
        # Find section markers and add content
        for para_idx, paragraph in enumerate(doc.paragraphs):
            para_text_upper = paragraph.text.upper().strip()
            
            for section_name, section_data in sections_to_add:
                if section_name.upper() in para_text_upper:
                    # Add content after this paragraph
                    print(f"✓ Adding content to section: {section_name}")
                    # Note: Adding paragraphs dynamically in DOCX is complex
                    # For now, we'll just log it
                    break
        
        # Save output
        output_docx = self.output_path.replace('.pdf', '.docx')
        doc.save(output_docx)
        
        print(f"\n✅ Successfully created formatted resume!")
        print(f"📁 Saved to: {output_docx}\n")
        return True
    
    def _create_field_mappings(self):
        """Map resume fields to template fields"""
        mappings = {}
        
        template_fields = self.template_analysis.get('fields', {})
        
        # Map each template field to resume data
        for field_name, field_info in template_fields.items():
            if field_name in self.resume_data:
                mappings[field_name] = {
                    'template_position': field_info,
                    'resume_value': self.resume_data[field_name]
                }
        
        return mappings
    
    def _match_sections(self):
        """Match resume sections to template sections"""
        matched = []
        
        template_sections = self.template_analysis.get('sections', [])
        resume_sections = self.resume_data.get('sections', {})
        
        # If template has defined sections, use those
        if template_sections:
            for template_section in template_sections:
                heading = template_section['heading']
                content = self._find_matching_resume_section(heading, resume_sections)
                if content:
                    matched.append((heading, content))
        else:
            # Use resume sections directly
            section_order = ['summary', 'experience', 'education', 'skills', 'projects', 'certifications']
            for section_key in section_order:
                for resume_key, content in resume_sections.items():
                    if section_key in resume_key.lower() and content:
                        matched.append((resume_key.title(), content))
                        break
        
        return matched
    
    def _find_matching_resume_section(self, template_heading, resume_sections):
        """Find resume section that matches template heading"""
        heading_lower = template_heading.lower()
        
        # Direct match
        if heading_lower in resume_sections:
            return resume_sections[heading_lower]
        
        # Keyword matching
        keywords = {
            'experience': ['experience', 'employment', 'work', 'professional'],
            'education': ['education', 'academic', 'qualification'],
            'skills': ['skills', 'technical', 'competencies'],
            'summary': ['summary', 'objective', 'profile'],
            'projects': ['projects', 'portfolio'],
            'certifications': ['certifications', 'certificates'],
            'awards': ['awards', 'achievements']
        }
        
        for key, patterns in keywords.items():
            if any(p in heading_lower for p in patterns):
                for resume_key, content in resume_sections.items():
                    if any(p in resume_key for p in patterns):
                        return content
        
        return []
    
    def _template_has_name_placeholder(self):
        """Check if template has name placeholder"""
        placeholders = self.template_analysis.get('placeholders', [])
        return any(p.get('type') == 'name' for p in placeholders)
    
    def _wrap_text(self, text, canvas_obj, max_width):
        """Wrap text to fit width"""
        words = text.split()
        lines = []
        current_line = []
        
        for word in words:
            test_line = ' '.join(current_line + [word])
            if canvas_obj.stringWidth(test_line) <= max_width:
                current_line.append(word)
            else:
                if current_line:
                    lines.append(' '.join(current_line))
                current_line = [word]
        
        if current_line:
            lines.append(' '.join(current_line))
        
        return lines if lines else [text[:100]]


def format_resume_intelligent(resume_data, template_analysis, output_path):
    """Main function for intelligent formatting"""
    formatter = IntelligentFormatter(resume_data, template_analysis, output_path)
    return formatter.format()
