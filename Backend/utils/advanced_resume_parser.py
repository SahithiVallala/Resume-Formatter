"""
Advanced Resume Parser
Extracts comprehensive information from resumes:
- Personal Details: Name, DOB, Email, Phone, Address, LinkedIn
- Experience: Company, Role, Duration, Responsibilities
- Education: Degree, Institution, Year, Grade
- Skills: Technical, Soft, Languages
- Projects, Certifications, Awards
"""

import pdfplumber
from docx import Document
import re
from datetime import datetime
from collections import defaultdict
import os

class ResumeParser:
    """Comprehensive resume parsing"""
    
    def __init__(self, file_path, file_type):
        self.file_path = file_path
        self.file_type = file_type
        self.raw_text = ""
        self.lines = []
        
    def parse(self):
        """Main parsing method"""
        print(f"\n{'='*70}")
        print(f"📋 PARSING RESUME: {self.file_path.split('/')[-1]}")
        print(f"{'='*70}\n")
        
        # Extract text
        if self.file_type == 'pdf':
            self.raw_text = self._extract_pdf_text()
        else:
            self.raw_text = self._extract_docx_text()
        
        self.lines = [line.strip() for line in self.raw_text.split('\n') if line.strip()]
        
        # Extract all information
        resume_data = {
            'name': self._extract_name(),
            'email': self._extract_email(),
            'phone': self._extract_phone(),
            'address': self._extract_address(),
            'linkedin': self._extract_linkedin(),
            'dob': self._extract_dob(),
            'summary': self._extract_summary(),
            'experience': self._extract_experience(),
            'education': self._extract_education(),
            'skills': self._extract_skills(),
            'projects': self._extract_projects(),
            'certifications': self._extract_certifications(),
            'awards': self._extract_awards(),
            'languages': self._extract_languages(),
            'sections': self._extract_sections(),
            'raw_text': self.raw_text
        }
        
        self._print_parsing_summary(resume_data)
        return resume_data
    
    def _extract_pdf_text(self):
        """Extract text from PDF"""
        try:
            with pdfplumber.open(self.file_path) as pdf:
                return '\n'.join([page.extract_text() or '' for page in pdf.pages])
        except Exception as e:
            print(f"❌ Error extracting PDF text: {e}")
            return ""
    
    def _extract_docx_text(self):
        """Extract text from DOCX (paragraphs, tables, headers/footers)"""
        try:
            doc = Document(self.file_path)
            texts = []

            # Main body paragraphs
            for para in doc.paragraphs:
                if para.text is not None:
                    texts.append(para.text)

            # Tables content
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if p.text is not None:
                                texts.append(p.text)

            # Headers and footers
            try:
                for section in doc.sections:
                    for p in section.header.paragraphs:
                        if p.text is not None:
                            texts.append(p.text)
                    for p in section.footer.paragraphs:
                        if p.text is not None:
                            texts.append(p.text)
            except Exception:
                pass

            # Join and normalize lines
            return '\n'.join([t for t in texts if t and t.strip()])
        except Exception as e:
            print(f"❌ Error extracting DOCX text: {e}")
            return ""
    
    def _extract_name(self):
        """Extract candidate name robustly using context around contact lines"""
        lines = self.lines

        # Helper: find first line index that contains email or phone
        email_idx = phone_idx = None
        email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b')
        phone_patterns = [
            r'(?:\+?\d{1,3}[-\s]?)?\(?\d{3}\)?[-\s]?\d{3}[-\s]?\d{4}',  # US formats
            r'(?:\+?91[-\s]?)?\d{10}',  # India +91XXXXXXXXXX or 10 digits
            r'\+?\d{1,3}[-\s]?\d{2,5}[-\s]?\d{5,10}',  # Intl generic
        ]
        phone_regexes = [re.compile(p) for p in phone_patterns]

        for idx, line in enumerate(lines[:30]):
            if email_idx is None and email_pattern.search(line):
                email_idx = idx
            if phone_idx is None and any(r.search(line) for r in phone_regexes):
                phone_idx = idx
            if email_idx is not None and phone_idx is not None:
                break

        # Try lines above the earliest contact line
        pivot = min([i for i in [email_idx, phone_idx] if i is not None], default=None)
        if pivot is not None:
            for i in range(max(0, pivot - 3), pivot + 1):
                line = lines[i].strip()
                if self._has_contact_info(line):
                    continue
                # Name heuristics: 2-4 words, mostly capitalized
                words = line.split()
                if 1 < len(words) <= 5 and sum(w[:1].isupper() for w in words) >= max(2, len(words) - 1):
                    return line

        # Fallback: scan first 10 non-contact lines for a likely name
        for line in lines[:10]:
            if self._has_contact_info(line):
                continue
            if len(line) < 3 or len(line) > 60:
                continue
            if re.match(r'^[A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z\-\.]+){1,3}$', line.strip()):
                return line.strip()

        # Fallback 2: derive from file name
        try:
            base = os.path.basename(self.file_path)
            stem = os.path.splitext(base)[0]
            # Remove common words and separators
            tokens = re.split(r'[\W_]+', stem)
            blacklist = {"resume", "cv", "profile", "updated", "final", "copy", "doc", "docx", "pdf"}
            tokens = [t for t in tokens if t and t.lower() not in blacklist]
            if 1 <= len(tokens) <= 4:
                name_guess = ' '.join(tokens)
                # Capitalize words
                return ' '.join(w[:1].upper() + w[1:] for w in name_guess.split())
        except Exception:
            pass

        return "Unknown Candidate"
    
    def _extract_email(self):
        """Extract email address"""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        match = re.search(email_pattern, self.raw_text)
        return match.group(0) if match else ""
    
    def _extract_phone(self):
        """Extract phone number (supports international formats)"""
        phone_patterns = [
            r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # US formats
            r'(?:\+?91[-.\s]?)?\d{10}',  # India
            r'\+?\d{1,3}[-.\s]?\d{2,5}[-.\s]?\d{5,10}',  # Intl generic
        ]

        for pattern in phone_patterns:
            match = re.search(pattern, self.raw_text)
            if match:
                return match.group(0)
        return ""
    
    def _extract_address(self):
        """Extract address"""
        # Look for lines with address keywords
        address_keywords = ['street', 'avenue', 'road', 'city', 'state', 'zip', 'postal']
        for line in self.lines:
            if any(keyword in line.lower() for keyword in address_keywords):
                return line
        return ""
    
    def _extract_linkedin(self):
        """Extract LinkedIn profile"""
        linkedin_pattern = r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+'
        match = re.search(linkedin_pattern, self.raw_text, re.IGNORECASE)
        return match.group(0) if match else ""
    
    def _extract_dob(self):
        """Extract date of birth"""
        dob_patterns = [
            r'(?:DOB|Date of Birth|Birth Date)[\s:]+(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})',
            r'(?:DOB|Date of Birth|Birth Date)[\s:]+(\w+ \d{1,2},? \d{4})',
        ]
        
        for pattern in dob_patterns:
            match = re.search(pattern, self.raw_text, re.IGNORECASE)
            if match:
                return match.group(1)
        return ""
    
    def _extract_summary(self):
        """Extract professional summary/objective"""
        summary_keywords = ['summary', 'objective', 'profile', 'about']
        
        for i, line in enumerate(self.lines):
            if any(keyword in line.lower() for keyword in summary_keywords):
                # Get next few lines as summary
                summary_lines = []
                for j in range(i+1, min(i+6, len(self.lines))):
                    if self._is_section_header(self.lines[j]):
                        break
                    summary_lines.append(self.lines[j])
                return ' '.join(summary_lines)
        return ""
    
    def _extract_experience(self):
        """Extract work experience details with robust pairing (dates + role/company)."""
        experiences = []
        section = self._find_section(['experience', 'work history', 'employment', 'professional experience'])
        if not section:
            return experiences

        # Normalize lines
        lines = [self._normalize_text(l) for l in section if l and l.strip()]
        i = 0
        while i < len(lines):
            line = lines[i]

            # Case A: Date line followed by role/company line
            if self._contains_date_range(line):
                duration = self._clean_years(line)
                # Find next descriptive line for role/company
                j = i + 1
                while j < len(lines) and (not lines[j] or self._is_section_header(lines[j])):
                    j += 1
                title_line = lines[j] if j < len(lines) else ''
                # Prefer company from the next line (strip bullet/location)
                company_line = self._strip_bullet(title_line)
                company, role = self._parse_company_role_line(company_line)
                # If parsing by delimiter failed, treat next line as company and role from date line
                if not company and company_line:
                    company = self._strip_location(company_line)
                # Derive role from the dated line by removing dates/tokens
                role_from_date = self._extract_role_from_dated_line(line)
                if role_from_date:
                    role = role_from_date
                company, role = self._strip_location(company), self._strip_location(role)

                exp = {
                    'company': company,
                    'role': role,
                    'title': title_line,
                    'duration': duration,
                    'details': []
                }

                # Collect following detail bullets until next date or header
                k = j + 1
                while k < len(lines):
                    if self._contains_date_range(lines[k]) or self._looks_like_company_or_role(lines[k]) or self._is_section_header(lines[k]):
                        break
                    exp['details'].append(lines[k])
                    k += 1
                experiences.append(exp)
                i = k
                continue

            # Case B: Role/company line followed by date line
            if self._looks_like_company_or_role(line):
                company, role = self._parse_company_role_line(line)
                company, role = self._strip_location(company), self._strip_location(role)
                # Look ahead for date
                j = i + 1
                duration = ''
                if j < len(lines) and self._contains_date_range(lines[j]):
                    duration = self._clean_years(lines[j])
                    j += 1
                exp = {
                    'company': company,
                    'role': role,
                    'title': line,
                    'duration': duration,
                    'details': []
                }
                # Collect details
                k = j
                while k < len(lines):
                    if self._contains_date_range(lines[k]) or self._looks_like_company_or_role(lines[k]) or self._is_section_header(lines[k]):
                        break
                    exp['details'].append(lines[k])
                    k += 1
                experiences.append(exp)
                i = k
                continue

            i += 1

        return experiences
    
    def _parse_company_role_line(self, line):
        """Parse a line to extract company and role"""
        # Common patterns:
        # "Company Name - Role Title"
        # "Role Title at Company Name"
        # "Company Name, Role Title"
        # "Role Title | Company Name"
        
        if ' - ' in line or ' – ' in line:
            # Split by dash
            parts = re.split(r'\s+[-–]\s+', line, maxsplit=1)
            if len(parts) == 2:
                # Usually "Company - Role" or "Role - Company"
                # Heuristic: if first part has common company indicators, it's company
                if any(word in parts[0].lower() for word in ['inc', 'corp', 'ltd', 'llc', 'company', 'technologies', 'systems']):
                    return parts[0].strip(), parts[1].strip()
                else:
                    # Assume first is role, second is company
                    return parts[1].strip(), parts[0].strip()
        
        elif ' at ' in line.lower():
            # "Role at Company"
            parts = re.split(r'\s+at\s+', line, flags=re.IGNORECASE, maxsplit=1)
            if len(parts) == 2:
                return parts[1].strip(), parts[0].strip()
        
        elif '|' in line:
            # "Role | Company"
            parts = line.split('|', maxsplit=1)
            if len(parts) == 2:
                return parts[1].strip(), parts[0].strip()
        
        elif ',' in line:
            # "Company, Role" or "Role, Company"
            parts = line.split(',', maxsplit=1)
            if len(parts) == 2:
                # Heuristic: if first part has company indicators
                if any(word in parts[0].lower() for word in ['inc', 'corp', 'ltd', 'llc', 'company', 'technologies', 'systems']):
                    return parts[0].strip(), parts[1].strip()
                else:
                    return parts[1].strip(), parts[0].strip()
        
        # If no pattern matched, assume entire line is company or role
        # Check if it looks more like a company name
        if any(word in line.lower() for word in ['inc', 'corp', 'ltd', 'llc', 'company', 'technologies', 'systems', 'solutions']):
            return line.strip(), ''
        else:
            # Assume it's a role
            return '', line.strip()
        
        return line.strip(), ''
    
    def _extract_education(self):
        """Extract education with pairing: degree line (with colon/institution) + year next line."""
        education = []
        section = self._find_section(['education', 'academic', 'qualification'])
        if not section:
            return education

        lines = [self._normalize_text(l) for l in section if l and l.strip()]
        i = 0
        while i < len(lines):
            line = lines[i]

            # Degree line typically contains ':' and an institution keyword on same line
            if ':' in line or any(k in line.lower() for k in ['university', 'college', 'school', 'institute', 'academy']):
                degree, institution = self._parse_degree_institution_line(line)
                degree = degree.strip()
                institution = self._strip_location(institution)

                # year can be on next line
                year = ''
                j = i + 1
                if j < len(lines) and self._contains_date_range(lines[j]):
                    year = self._clean_years(lines[j])
                    i = j + 1
                else:
                    # Try to find years on same line
                    year = self._clean_years(line)
                    i += 1

                edu = {
                    'degree': degree,
                    'institution': institution,
                    'year': year,
                    'details': []
                }
                education.append(edu)
                continue

            i += 1

        return education

    # ---------------------- Helpers ----------------------
    def _normalize_text(self, s):
        # Remove odd unicode like 'ï¼' and zero-width, collapse spaces
        if not s:
            return ''
        s = s.replace('\u200b', ' ').replace('ï¼', ' ').replace('\ufeff', ' ')
        s = re.sub(r'\s+', ' ', s).strip()
        return s

    def _strip_location(self, s):
        if not s:
            return s
        # Remove trailing city/state fragments after comma
        s = re.sub(r',[^,]*\b(?:city|state|india|usa|uk)\b.*$', '', s, flags=re.IGNORECASE)
        return s.strip()

    def _clean_years(self, s):
        if not s:
            return ''
        s = self._normalize_text(s)
        s = re.sub(r'\s*(to|–|—|-)\s*', '-', s, flags=re.IGNORECASE)
        s = re.sub(r'\b(current|present)\b', str(datetime.now().year), s, flags=re.IGNORECASE)
        years = re.findall(r'\b(?:19|20)\d{2}\b', s)
        if len(years) >= 2:
            return f"{years[0]}-{years[-1]}"
        elif len(years) == 1:
            return years[0]
        return ''
    
    def _parse_degree_institution_line(self, line):
        """Parse a line to extract degree and institution"""
        # Common patterns:
        # "Bachelor of Science, MIT"
        # "Master's in Computer Science - Stanford University"
        # "MBA | Harvard Business School"
        s = self._normalize_text(line)
        s = self._strip_bullet(s)
        # Prefer splitting at an institution keyword boundary
        m = re.search(r'(university|college|school|institute|academy)\b.*', s, flags=re.IGNORECASE)
        if m:
            degree = s[:m.start()].strip(' ,;:-')
            institution = s[m.start():].strip()
            # Remove trailing location fragments
            institution = self._strip_location(institution)
            return degree, institution
        
        if ',' in s:
            parts = s.split(',', maxsplit=1)
            return parts[0].strip(), parts[1].strip()
        
        elif ' - ' in s or ' – ' in s:
            parts = re.split(r'\s+[-–]\s+', s, maxsplit=1)
            if len(parts) == 2:
                return parts[0].strip(), parts[1].strip()
        
        elif '|' in s:
            parts = s.split('|', maxsplit=1)
            if len(parts) == 2:
                return parts[0].strip(), parts[1].strip()
        
        elif ' from ' in s.lower():
            parts = re.split(r'\s+from\s+', s, flags=re.IGNORECASE, maxsplit=1)
            if len(parts) == 2:
                return parts[0].strip(), parts[1].strip()
        
        # If no pattern, return entire line as degree
        return s.strip(), ''

    def _strip_bullet(self, s):
        if not s:
            return s
        return re.sub(r'^[\s•\-–—*●]+', '', s)

    def _extract_role_from_dated_line(self, s):
        """Remove date expressions from a dated line to recover the job title/role text."""
        if not s:
            return ''
        t = self._strip_bullet(self._normalize_text(s))
        # Remove month names + year e.g., 'Aug 2007', 'July 2025'
        t = re.sub(r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-zA-Z]*\s+(?:19|20)\d{2}\b', '', t, flags=re.IGNORECASE)
        # Remove standalone years and connectors
        t = re.sub(r'\b(?:19|20)\d{2}\b', '', t)
        t = re.sub(r'\b(to|–|—|-)\b', '', t, flags=re.IGNORECASE)
        t = re.sub(r'\b(current|present)\b', '', t, flags=re.IGNORECASE)
        # Collapse spaces and trim
        t = re.sub(r'\s+', ' ', t).strip(' ,;:- ')
        return t
    
    def _extract_skills(self):
        """Extract skills"""
        skills = []
        skills_section = self._find_section(['skills', 'technical skills', 'competencies', 'expertise'])
        
        if skills_section:
            for line in skills_section:
                # Split by common delimiters
                if ',' in line or '|' in line or '•' in line:
                    parts = re.split(r'[,|•]', line)
                    skills.extend([p.strip() for p in parts if p.strip()])
                elif line and not self._is_section_header(line):
                    skills.append(line)
        
        return skills
    
    def _extract_projects(self):
        """Extract projects"""
        projects = []
        projects_section = self._find_section(['projects', 'key projects', 'project work'])
        
        if projects_section:
            current_project = {}
            for line in projects_section:
                if line and not line.startswith('•') and not self._is_section_header(line):
                    if current_project:
                        projects.append(current_project)
                    current_project = {'name': line, 'details': []}
                elif line:
                    current_project.setdefault('details', []).append(line)
            
            if current_project:
                projects.append(current_project)
        
        return projects
    
    def _extract_certifications(self):
        """Extract certifications"""
        certifications = []
        cert_section = self._find_section(['certifications', 'certificates', 'licenses'])
        
        if cert_section:
            for line in cert_section:
                if line and not self._is_section_header(line):
                    certifications.append(line)
        
        return certifications
    
    def _extract_awards(self):
        """Extract awards and achievements"""
        awards = []
        awards_section = self._find_section(['awards', 'achievements', 'honors', 'recognition'])
        
        if awards_section:
            for line in awards_section:
                if line and not self._is_section_header(line):
                    awards.append(line)
        
        return awards
    
    def _extract_languages(self):
        """Extract languages"""
        languages = []
        lang_section = self._find_section(['languages', 'language proficiency'])
        
        if lang_section:
            for line in lang_section:
                if line and not self._is_section_header(line):
                    languages.append(line)
        
        return languages
    
    def _extract_sections(self):
        """Extract all sections for flexible formatting"""
        sections = defaultdict(list)
        current_section = None
        
        for line in self.lines:
            if self._is_section_header(line):
                current_section = line.lower().strip()
            elif current_section and line:
                sections[current_section].append(line)
        
        return dict(sections)
    
    def _find_section(self, keywords):
        """Find section by keywords"""
        section_lines = []
        in_section = False
        
        for line in self.lines:
            if any(keyword in line.lower() for keyword in keywords):
                in_section = True
                continue
            
            if in_section:
                if self._is_section_header(line):
                    break
                section_lines.append(line)
        
        return section_lines
    
    def _is_section_header(self, line):
        """Check if line is a section header"""
        if len(line) > 50:
            return False
        
        section_keywords = [
            'experience', 'education', 'skills', 'summary', 'objective',
            'projects', 'certifications', 'awards', 'languages', 'profile',
            'work history', 'employment', 'qualifications', 'achievements'
        ]
        
        line_lower = line.lower().strip()
        return any(keyword == line_lower or line_lower.startswith(keyword) for keyword in section_keywords)
    
    def _has_contact_info(self, text):
        """Check if text contains contact information"""
        return bool(re.search(r'@|http|linkedin|\d{3}[-.\s]\d{3}', text, re.IGNORECASE))
    
    def _looks_like_company_or_role(self, line):
        """Check if line looks like company name or job title"""
        # Usually title case, not too long, no bullets
        return (line.istitle() or line.isupper()) and len(line) < 100 and not line.startswith('•')
    
    def _contains_date_range(self, line):
        """Check if line contains date range"""
        date_patterns = [
            r'\b(19|20)\d{2}\b.*\b(19|20)\d{2}\b',  # 2020 - 2023
            r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+\d{4}',  # Jan 2020
            r'\d{1,2}/\d{4}',  # 01/2020
        ]
        return any(re.search(pattern, line, re.IGNORECASE) for pattern in date_patterns)
    
    def _print_parsing_summary(self, data):
        """Print parsing summary"""
        print(f"👤 Name: {data['name']}")
        print(f"📧 Email: {data['email']}")
        print(f"📱 Phone: {data['phone']}")
        print(f"🔗 LinkedIn: {data['linkedin'][:50] if data['linkedin'] else 'Not found'}")
        print(f"📅 DOB: {data['dob']}")
        print(f"💼 Experience Entries: {len(data['experience'])}")
        print(f"🎓 Education Entries: {len(data['education'])}")
        print(f"🛠️  Skills: {len(data['skills'])}")
        print(f"📂 Projects: {len(data['projects'])}")
        print(f"🏆 Certifications: {len(data['certifications'])}")
        print(f"🏅 Awards: {len(data['awards'])}")
        print(f"🌐 Languages: {len(data['languages'])}")
        print(f"\n{'='*70}\n")


def parse_resume(file_path, file_type):
    """Main function to parse resume"""
    parser = ResumeParser(file_path, file_type)
    return parser.parse()
